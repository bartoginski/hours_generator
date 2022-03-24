# Import docx NOT python-docx
import docx
from date import get_year, get_month


def generate_doc(filename, contract_date, programmer_name, hours, total, output_path="./"):
    month = get_month(contract_date)
    year = get_year(contract_date)
    doc = docx.Document()
    doc.add_heading('Rejestr godzin relizacji zlecenia', 1)
    doc.add_paragraph(
        f'''Rozliczenie liczby godzin wykonanych usług do umowy zlecenie z dnia {contract_date}
w {month} {year}''')
    doc.add_paragraph(f'Zleceniobiorca: {programmer_name}')

    data = []
    for i, time in enumerate(hours):
        data.append((str(i + 1), str(time), '', ''))

    table = doc.add_table(rows=1, cols=4)

    row = table.rows[0].cells
    row[0].text = 'Dzień miesiąca'
    row[1].text = 'Liczba godzin realizacji zlecenia'
    row[2].text = 'Podpis zleceniobiorcy'
    row[3].text = 'Podpis zleceniodawcy'

    for i, hours, empty_1, empty_2 in data:
        row = table.add_row().cells
        row[0].text = i
        row[1].text = hours
        row[2].text = empty_1
        row[3].text = empty_2

    row = table.add_row().cells
    row[0].text = 'Łącznie'
    row[1].text = str(total)
    row[2].text = ''
    row[3].text = ''

    # table style
    table.style = 'Table Grid'
    doc.save(f"{output_path}{filename}.docx")
